from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_bold_then_text(bold_text, normal_text):
    p = doc.add_paragraph()
    r1 = p.add_run(bold_text)
    r1.bold = True
    p.add_run(normal_text)
    return p

# TITLE
p = doc.add_heading('KARIEGA CHAMPS PUBLIC ART', level=0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
r = p.add_run('NAC Annual Project Funding 2026-2027 — Final Refined Submission')
r.bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('Organisation: APBC Investments | Discipline: Visual Arts')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# SECTION 1
doc.add_heading('SECTION 1: Project Impact Type', level=1)
doc.add_paragraph('Select: Economic').runs[0].bold = True
doc.add_heading('1.1 Describe the benefit of the project (~100 words)', level=2)
doc.add_paragraph(
    'The Kariega Champs Public Art project invests R152,000 directly in employee stipends, creating 24 paid positions across fine art, project management, event coordination, and audio-visual production in Kariega. Creative roles attract stipends of R4,000\u2013R5,000 per month; interns earn R2,000 per month \u2014 meaningful income for township households. By prioritising women (15 of 24 roles), youth, and individuals with disabilities (3 reserved positions), nearly half the R307,000 NAC request flows directly into historically disadvantaged communities. Combined with R180,000 in confirmed in-kind contributions from MpumaKapa TV and HD Sounds, the total project value of R487,000 demonstrates strong financial leverage and tangible community economic investment.'
)

# SECTION 2
doc.add_heading('SECTION 2: Focus Areas', level=1)
p = doc.add_paragraph()
p.add_run('Select all four:').bold = True
doc.add_paragraph('Social Cohesion and Nation Building', style='List Bullet')
doc.add_paragraph('Marginalised and Indigenous Arts', style='List Bullet')
doc.add_paragraph('Addressing Social Ills', style='List Bullet')
doc.add_paragraph('Supporting Vulnerable Groups', style='List Bullet')

doc.add_heading('2.1 Comment on focus areas selected (~500 words)', level=2)

add_bold_then_text(
    'Social Cohesion and Nation Building \u2014 ',
    'The Kariega Champs Public Art project fosters social cohesion by creating three large-scale murals in public spaces that celebrate a shared heritage cutting across racial, generational, and socio-economic lines. Mzukisi Sikali, Vuyani Nene, and Zolani Petelo are heroes whose achievements belong to all of Kariega, not to a single demographic. By placing their stories on public walls in both KwaNobuhle and central Kariega, the project creates gathering points where residents from different backgrounds reconnect with a collective identity rooted in resilience and excellence. The community co-creation process \u2014 including a public forum to gather input on mural themes and locations \u2014 ensures that diverse voices shape the final artworks, building trust between residents, local leaders, and the creative team. In a town marked by the 1985 anti-apartheid protests, celebrating local champions through collaborative public art contributes meaningfully to nation building by honouring struggle history alongside sporting triumph.'
)

add_bold_then_text(
    'Marginalised and Indigenous Arts \u2014 ',
    'The mural designs deliberately integrate graphic elements inspired by isiXhosa visual traditions alongside large-scale realistic portraiture, bridging Indigenous cultural expression with contemporary public art practice. This fusion ensures that the murals are not simply biographical tributes but living expressions of the cultural identity of KwaNobuhle and surrounding townships. The project\u2019s mentorship model places five emerging artists from marginalised communities at the centre of the creative process, working under the lead artist through every stage from concept sketching to final execution. This elevates artists who would otherwise lack access to professional mural commissions, transferring technical skills in large-format painting, surface preparation, and community-engaged design. By prioritising local, township-based talent and embedding Indigenous visual language into the artwork, the project challenges the exclusion of marginalised practitioners from high-profile public art opportunities.'
)

add_bold_then_text(
    'Addressing Social Ills \u2014 ',
    'Kariega faces persistent unemployment, particularly among youth, and the associated social challenges of substance abuse, crime, and disengagement. The project responds directly by generating 24 paid positions across fine art, project management, audio-visual production, and event coordination. These are structured roles offering mentorship and real work experience that build employable skills transferable beyond the project\u2019s lifespan. The emphasis on hiring from within Kariega ensures that income circulates locally, strengthening household stability and reducing economic vulnerability. By channelling creative energy into constructive community outcomes, the project provides a tangible alternative to the negative influences that thrive where opportunity is absent. The public unveiling events and social media workshops further equip participants with digital literacy and event management competencies, broadening their future prospects.'
)

add_bold_then_text(
    'Supporting Vulnerable Groups \u2014 ',
    'APBC Investments, the applicant organisation, is 100% women-owned and women-led, and gender parity is built into the project\u2019s staffing plan: 15 of the 24 positions are allocated to women. Three positions are reserved for individuals living with disabilities. Mural sites, unveiling events, and workshops will be designed for universal accessibility, including wheelchair-friendly pathways and sign language interpretation at public gatherings. Recruitment will specifically target women and youth from underrepresented racial backgrounds within Kariega, prioritising those from low-income households. These are not afterthoughts but core design principles that ensure vulnerable groups participate as creators and beneficiaries, not merely as audiences.'
)

# SECTION 3
doc.add_heading('SECTION 3: Project Impact Categories', level=1)
p = doc.add_paragraph()
p.add_run('Select all three:').bold = True
doc.add_paragraph('Youth', style='List Bullet')
doc.add_paragraph('Women', style='List Bullet')
doc.add_paragraph('People living with disability', style='List Bullet')

doc.add_heading('3.1 Comment on project impact categories (~500 words)', level=2)

add_bold_then_text(
    'Youth \u2014 ',
    'The project allocates the majority of its 24 positions to young people aged 18\u201335, with stipends ranging from R2,000 per month for Fine Art Interns to R4,000\u2013R6,500 per month for skilled roles such as Fine Artists, Social Media Officers, and Project Managers. This represents R152,000 in direct income channelled into Kariega households over the project\u2019s duration. Youth will fill paid creative roles (5 artists at R4,000/month, 5 interns at R2,000/month), project management positions (R3,000\u2013R6,500/month), and technical support roles. Each participant will be paired with an experienced mentor who will guide them through structured milestones covering creative technique, professional conduct, and career planning. Youth from diverse racial backgrounds across Kariega\u2019s townships will be actively recruited, advancing the NAC\u2019s transformation imperatives. Beyond stipends, the R21,000 Community Visioning Workshop provides hands-on training in community-engaged design, while the social media component equips youth with digital marketing and content creation skills with direct commercial application. Post-project, participants receive certificates of completion and professional references, formalising their experience for future employment.'
)

add_bold_then_text(
    'Women \u2014 ',
    'Gender parity is a foundational principle of this project. APBC Investments is 100% women-owned and women-led, with Project Director Aphelele Beyi leading both creative vision and operational management. Of the 24 positions, 15 are allocated to women \u2014 representing over R90,000 in stipends directed to women from Kariega\u2019s townships. Women will occupy roles at every level: from Fine Art Interns (R2,000/month) to the Creative Director (R5,000/month), Marketing Manager (R4,000/month), and Finance Manager (R5,000/month), ensuring women hold leadership and skilled positions rather than being concentrated in lower-paid support functions. Recruitment will prioritise women from underrepresented racial backgrounds and low-income households. Women participants will gain practical skills in arts project planning, event coordination, technical production, and financial management \u2014 competencies that translate directly into employability beyond this project. Mentorships led by experienced female professionals within the core team will support personal growth and sector-specific networking. Through paid employment at meaningful stipend levels, women gain financial independence and contribute to household stability, modelling women\u2019s economic participation for the broader Kariega community.'
)

add_bold_then_text(
    'People with Disabilities \u2014 ',
    'Three of the 24 positions are reserved for individuals living with disabilities, ensuring their active participation as creators and contributors, not merely as passive beneficiaries. The project will partner with local disability organisations in Kariega to identify candidates and ensure appropriate support structures are in place. All project activities \u2014 workshops, community forums, mural site visits, and unveiling events \u2014 will be designed for universal accessibility: wheelchair-friendly pathways, braille and large-print signage, and sign language interpretation at public gatherings. Inclusive creative workshops will accommodate a range of physical and sensory abilities, ensuring every participant\u2019s artistic potential is developed and valued. The visible participation of people with disabilities in a high-profile public art project challenges stigma, raises community awareness, and models inclusive practice for the broader arts sector. Outcomes for this group will be tracked through participation data, accessibility audits, and qualitative feedback from participants and their support networks.'
)

doc.add_paragraph('Across all three categories, impact will be measured through employment data, skills acquisition assessments, and participant feedback surveys, ensuring transparent reporting to the NAC and accountability to the communities the project serves.')

# SECTION 4
doc.add_heading('SECTION 4: Briefly describe your project (~1,000 words)', level=1)

doc.add_heading('INTRODUCTION', level=3)
doc.add_paragraph('The Kariega Champs Public Art project is a community-driven visual arts initiative that will create three large-scale murals in public spaces across Kariega, Eastern Cape, celebrating the town\u2019s rich boxing heritage through the stories of three local champions: Mzukisi \u201cLacier\u201d Sikali, Vuyani \u201cWonderboy\u201d Nene, and Zolani Petelo. More than a commemoration of sporting achievement, the project is designed as a vehicle for community development \u2014 generating 24 employment opportunities, transferring creative skills to emerging artists, and fostering social cohesion in a historically disadvantaged area within Nelson Mandela Bay Municipality.')

doc.add_heading('CONTEXT AND COMMUNITY NEED', level=3)
doc.add_paragraph('Kariega, formerly Uitenhage, is an industrial town with deep historical significance. Its role in the 1985 anti-apartheid protests \u2014 when police killed unarmed mourners \u2014 is seared into the national memory. Today, the town\u2019s townships face persistent unemployment, particularly among youth, and a lack of visible cultural infrastructure that celebrates local identity. While Kariega has produced nationally and internationally recognised athletes across multiple sports, these achievements remain largely uncelebrated in the town\u2019s public spaces. The project addresses that gap by placing the stories of local heroes where they belong: on the walls of their community, creating permanent markers of pride and possibility for current and future generations. This directly serves the NAC\u2019s mandate to promote arts application in communities and to address historical infrastructure imbalances.')

doc.add_heading('THE THREE MURALS', level=3)
doc.add_paragraph('Each mural will be a large-scale, realistic portrait combined with graphic elements inspired by isiXhosa visual traditions, creating artworks that are simultaneously contemporary and rooted in local cultural identity.')

add_bold_then_text('Mzukisi \u201cLacier\u201d Sikali \u2014 ', 'Born in KwaNobuhle, Sikali became a three-division world boxing champion, conquering the junior flyweight, junior bantamweight, and flyweight divisions. His professional record of 29 victories included title wins in Bangkok, Italy, and a celebrated homecoming victory for the IBO flyweight title. Tragically killed in 2005 in KwaNobuhle, his story embodies both the extraordinary potential and the urgent need for opportunity within Kariega\u2019s communities.')

add_bold_then_text('Vuyani \u201cWonderboy\u201d Nene \u2014 ', 'A revered figure from the 1980s and 1990s boxing era, Nene\u2019s legacy rests on his exceptional technical skill and four unprecedented victories over world champion Baby Jake Matlala. Across 63 professional fights (46 wins, 11 losses, 6 draws), his consistency and tactical brilliance earned him recognition as one of South Africa\u2019s most respected boxers.')

add_bold_then_text('Zolani Petelo \u2014 ', 'Petelo captured the IBF mini-flyweight world title in 1997 with a stunning fourth-round knockout of the previously undefeated Ratanapol Sor Vorapin, ending Sor Vorapin\u2019s streak of 20 consecutive title fight victories. He successfully defended his belt against international contenders and later challenged the legendary Ricardo Lopez at Madison Square Garden, demonstrating the courage that defined his 22-win career.')

doc.add_heading('HOW THE PROJECT WILL BE EXECUTED', level=3)
doc.add_paragraph('The project follows a structured, phased approach over three months (January\u2013March 2026):')

add_bold_then_text('Phase 1: Community Engagement and Design (January 2026) \u2014 ', 'The project begins with stakeholder identification and community consultation. A public Kariega Community Forum will gather input on mural themes, locations, and narratives, ensuring that the final artworks reflect community values and aspirations. Artists will sketch initial concepts based on this input and refine designs through iterative feedback. A social media workshop will be launched to train participants in documenting the creative process.')

add_bold_then_text('Phase 2: Design Finalisation and Site Preparation (February 2026) \u2014 ', 'Mural designs will be finalised incorporating community feedback and isiXhosa graphic elements. Materials will be procured from local and regional suppliers, sites prepared (cleaning, priming, scaffolding), and permits secured. Media engagement begins with press releases, artist interviews, and social media content to build public anticipation and document the process.')

add_bold_then_text('Phase 3: Mural Creation and Unveilings (February\u2013March 2026) \u2014 ', 'The three murals will be painted sequentially, each by a team that includes the lead artist and five emerging artists working under a structured mentorship model. Each completed mural will be unveiled at a community celebration event featuring cultural performances, heritage discussions, and public engagement activities. MpumaKapa TV will produce broadcast-quality documentary content throughout this phase, extending the project\u2019s reach to provincial and national audiences.')

doc.add_heading('TEAM AND EMPLOYMENT', level=3)
doc.add_paragraph('The project will employ 24 individuals from Kariega, with positions spanning creative roles (15 positions: lead artist, emerging artists, interns), project management (6 positions: project director, manager, administration, marketing), and technical support (3 positions: audio-visual production, documentation). Recruitment prioritises women (15 of 24 positions), youth, and individuals with disabilities (3 reserved positions). APBC Investments, the applicant organisation, is 100% women-owned and women-led, with a core team comprising Project Director Aphelele Beyi, Artistic Director Noxolo Yekela, Project Manager Sinethemba Gayiza, Administration Officer Sonele Dyan, and Finance Manager Thando Cuntu. A structured mentorship programme pairs each emerging artist with an experienced professional, ensuring transfer of technical skills in large-format mural painting, surface preparation, community-engaged design, and project management.')

doc.add_heading('COMMUNITY IMPACT AND BUDGET', level=3)
doc.add_paragraph('The murals will become permanent cultural landmarks in Kariega, creating visible symbols of local pride in areas that currently lack cultural infrastructure. By combining public art creation with employment, skills transfer, and community engagement, the project delivers measurable social and economic outcomes for one of the Eastern Cape\u2019s most historically significant communities. The total project budget is R480,000, comprising R300,000 requested from the NAC and R180,000 in confirmed in-kind contributions from MpumaKapa TV (R100,000) and HD Sounds (R80,000). Additional funding applications have been submitted to Nelson Mandela Bay Municipality and ECPACC. Administrative expenses total R4,000 (1.3% of the budget), well within the NAC\u2019s 15% threshold.')

# SECTION 5
doc.add_heading('SECTION 5: Access to markets (~500 words)', level=1)
doc.add_paragraph('The Kariega Champs Public Art project employs a multi-channel strategy to reach diverse audiences locally, regionally, and nationally, ensuring the project\u2019s cultural impact extends well beyond the physical mural sites.')

add_bold_then_text('Permanent Public Access \u2014 ', 'The murals themselves are the primary audience access point. Located in public spaces across Kariega, they will be visible to thousands of residents and visitors daily, requiring no ticket, gallery visit, or digital access. Three community unveiling events will draw residents from KwaNobuhle and surrounding areas, combining mural reveals with cultural performances, panel discussions on local heritage, and interactive workshops. These events will be free and open to the public, removing economic barriers to participation. Partnerships with local schools and community organisations will ensure awareness reaches diverse demographics, including youth who may not typically engage with visual arts.')

add_bold_then_text('Media and Broadcast Partnership \u2014 ', 'MpumaKapa TV, a confirmed in-kind partner contributing R100,000, will produce broadcast-quality documentary content covering the mural creation process from concept to completion. This content will be distributed through MpumaKapa TV\u2019s regional broadcast network, reaching audiences across the Eastern Cape and beyond. The partnership extends the project\u2019s audience far beyond those who can physically visit the murals, positioning Kariega\u2019s boxing heritage within a broader cultural narrative accessible to provincial and national viewers.')

add_bold_then_text('Digital and Social Media Strategy \u2014 ', 'A dedicated social media campaign will run throughout the project using Facebook, Instagram, and TikTok to share real-time updates on mural progress, artist interviews, community reactions, and behind-the-scenes content. A social media workshop embedded within the project will train participating youth in content creation, digital storytelling, and audience engagement, ensuring the campaign is sustained by local voices rather than an external marketing team. This approach builds an online community around the project, attracting followers from the South African arts sector, boxing enthusiasts, and diaspora communities interested in Eastern Cape heritage.')

add_bold_then_text('Regional and National Arts Networks \u2014 ', 'The project will be profiled through arts sector networks, including submissions to the NAC\u2019s own communications channels, regional arts publications, and Eastern Cape cultural tourism platforms. The murals will be positioned as cultural landmarks on Nelson Mandela Bay\u2019s tourism routes, linking Kariega\u2019s boxing heritage to the broader tourism offering of Gqeberha and the surrounding region. Informational plaques at each mural site will provide historical context and QR codes linking to digital content, serving both local residents and visiting tourists.')

add_bold_then_text('Community and Educational Outreach \u2014 ', 'Partnerships with local schools will integrate the murals into educational programmes on history, visual arts, and community identity. Guided mural tours will be developed for school groups and community organisations, creating ongoing engagement opportunities beyond the project\u2019s timeline. These tours can be facilitated by project alumni \u2014 youth participants who gained skills during the project \u2014 creating a sustainable link between the artworks and their audience while providing continued employment for trained participants.')

doc.add_paragraph('Through this combination of permanent public art, broadcast media, digital storytelling, arts sector networking, and educational partnerships, the project ensures that Kariega\u2019s boxing heritage and the stories of Sikali, Nene, and Petelo reach audiences across multiple platforms and demographics.')

# SECTION 6
doc.add_heading('SECTION 6: Development and employment of youth (~500 words)', level=1)
doc.add_paragraph('The Kariega Champs Public Art project places youth development and employment at the centre of its design, recognising that meaningful economic participation is the most direct route to transforming the prospects of young people in historically disadvantaged communities.')

add_bold_then_text('Direct Employment \u2014 ', 'Of the 24 positions created by the project, the majority are allocated to young people aged 18\u201335. These are not volunteer or token roles but paid placements across multiple disciplines: mural painting, project management, event coordination, administration, audio-visual production, and digital content creation. Each role is structured with clear deliverables, timelines, and performance expectations, providing participants with genuine professional experience. Stipends are set at rates that reflect the value of the work and provide meaningful income for participants and their households.')

add_bold_then_text('Structured Mentorship \u2014 ', 'Every youth participant will be paired with an experienced professional from the core project team. The lead artist will mentor five emerging visual artists through the full mural creation process \u2014 from concept development and community consultation through to surface preparation, large-format painting techniques, and final detailing. Project management mentors will guide participants in planning, scheduling, budgeting, and stakeholder communication. These mentorships are structured programmes with weekly check-ins, milestone reviews, and feedback sessions designed to accelerate skills acquisition and professional growth.')

add_bold_then_text('Skills Transfer \u2014 ', 'The project deliberately builds competencies that extend beyond mural painting. Participants will gain skills in large-format visual arts techniques (surface preparation, scaling, paint application, weather-resistant finishing), community-engaged design (consultation processes, incorporating public feedback into creative work), digital marketing and content creation (social media management, photography, video documentation), event management (planning and executing public unveiling ceremonies), and project administration (scheduling, procurement, reporting). This breadth ensures that youth participants leave the project with a diverse portfolio of employable skills applicable across the creative industries and beyond.')

add_bold_then_text('Career Pathways \u2014 ', 'Upon completion, each participant will receive a certificate of completion and professional references from the core team. The project will host a networking session connecting participants with arts sector professionals, gallery owners, and cultural organisations in the Nelson Mandela Bay region. This creates tangible post-project pathways rather than leaving participants without follow-on opportunities. Participants who demonstrate strong aptitude will be recommended for future APBC projects and partner organisation opportunities.')

add_bold_then_text('Addressing Youth Unemployment \u2014 ', 'Kariega\u2019s youth unemployment rate reflects the broader Eastern Cape crisis, where young people in townships face limited access to formal employment, professional networks, and skills development. The project responds directly by bringing paid, structured work opportunities into the community rather than requiring youth to travel to distant urban centres for employment. By hiring locally, the project keeps income within Kariega and demonstrates that the creative industries can be a viable economic pathway for township youth.')

add_bold_then_text('Diversity and Inclusion \u2014 ', 'Recruitment will intentionally include youth from diverse racial backgrounds to reflect Kariega\u2019s multicultural demographics, advancing the NAC\u2019s redress and transformation imperatives. Youth from low-income households and those without previous formal employment will be prioritised, ensuring the project reaches those who stand to benefit most and who face the greatest barriers to economic participation.')

# SECTION 7
doc.add_heading('SECTION 7: Project sustainability (~500 words)', level=1)
doc.add_paragraph('The Kariega Champs Public Art project is designed to deliver lasting impact that extends well beyond the three-month implementation period and the NAC grant cycle. Sustainability is embedded in the project\u2019s structure through five interconnected strategies.')

add_bold_then_text('Permanent Public Art Infrastructure \u2014 ', 'Unlike performance-based or event-driven arts projects, the murals are permanent installations. Once completed, they will remain as cultural landmarks in Kariega for years to come, requiring minimal ongoing investment. Weather-resistant materials and UV-protective finishes will ensure long-term durability. The murals will continue to generate community pride, cultural engagement, and tourism interest long after the project concludes, delivering ongoing returns on the initial investment without recurring costs.')

add_bold_then_text('Skills Legacy \u2014 ', 'The mentorship model at the heart of the project creates a multiplier effect. The five emerging artists mentored during the mural creation process will carry forward professional competencies in large-format painting, community-engaged design, and project management. These individuals become a local talent pool capable of leading future public art commissions, community projects, and creative industry ventures in Kariega and the broader Eastern Cape. Similarly, youth participants trained in digital content creation, event management, and administration will carry these skills into future employment. The project\u2019s investment in human capital is its most enduring legacy.')

add_bold_then_text('Community Ownership \u2014 ', 'The community co-creation process \u2014 from the initial public forum through design feedback to the unveiling celebrations \u2014 builds a sense of collective ownership over the murals. Residents who contributed to the design and witnessed the creation process become advocates for the murals\u2019 preservation and maintenance. APBC will work with local ward councillors and community organisations to establish custodianship arrangements, ensuring that the mural sites are kept clean and accessible over the long term.')

add_bold_then_text('Partnership Continuity \u2014 ', 'The relationships built with MpumaKapa TV, HD Sounds, Nelson Mandela Bay Municipality, and ECPACC during this project will serve as the foundation for future collaborations. The documentary content produced by MpumaKapa TV will continue to be broadcast and shared after project completion, sustaining public interest. APBC will leverage this project\u2019s track record and documented outcomes to pursue follow-on funding for an expanded mural programme covering additional Kariega sporting and cultural figures, building on established partnerships rather than starting from scratch.')

add_bold_then_text('Revenue Potential and Institutional Growth \u2014 ', 'The mural trail concept has the potential to generate modest income through guided cultural tours, mural merchandise (prints, postcards), and partnerships with Nelson Mandela Bay tourism operators. While these revenue streams will take time to develop, they represent a pathway toward financial self-sufficiency for the maintenance and expansion of Kariega\u2019s public art assets. Additionally, this project strengthens APBC Investments\u2019 organisational capacity as a delivery agent for community arts projects. As a young, women-owned organisation, successfully executing a project of this scale \u2014 with transparent financial management, documented outcomes, and positive community impact \u2014 builds APBC\u2019s credibility and positions it for larger commissions and funding opportunities in future cycles.')

# SECTION 8
doc.add_heading('SECTION 8: Sustaining and attracting new partnerships (~500 words)', level=1)
doc.add_paragraph('The Kariega Champs Public Art project is built on a foundation of existing partnerships and a deliberate strategy to attract new collaborations that sustain the project\u2019s impact and expand its reach.')

doc.add_heading('Confirmed Partners', level=3)
add_bold_then_text('MpumaKapa TV (in-kind contribution: R100,000) \u2014 ', 'MpumaKapa TV is a regional broadcast media company that will provide professional audio-visual coverage of the entire mural creation process, from community consultation through to the final unveilings. This partnership delivers broadcast-quality documentary content that extends the project\u2019s audience beyond Kariega to provincial and national viewers. MpumaKapa TV\u2019s involvement strengthens the project\u2019s media profiling, generating public interest and validating the initiative within the broader arts and culture sector.')

add_bold_then_text('HD Sounds (in-kind contribution: R80,000) \u2014 ', 'HD Sounds will provide sound and stage equipment for the three community unveiling events. This partnership ensures that each unveiling is a professionally produced celebration that draws community attendance and media coverage, reinforcing the project\u2019s public engagement objectives and creating memorable community experiences.')

doc.add_heading('Funding Partners Under Development', level=3)
add_bold_then_text('Nelson Mandela Bay Municipality \u2014 ', 'APBC has submitted a funding application to the municipality (application date: 30 June 2025). Municipal support would provide additional financial resources and signal local government endorsement of the project, strengthening its credibility and opening doors to further institutional partnerships within the metro.')

add_bold_then_text('Eastern Cape Provincial Arts and Culture Council (ECPACC) \u2014 ', 'A funding application has been submitted (application date: 15 August 2025). Provincial support would position the project within the Eastern Cape\u2019s broader arts development strategy, connecting APBC to provincial networks and future commissioning opportunities.')

doc.add_heading('Community and Sector Partnerships', level=3)
doc.add_paragraph('Local community organisations, ward councillors, and school leadership in Kariega will be engaged as stakeholders from the project\u2019s inception. Their involvement in the community forum and design consultation process ensures that the project reflects local values and builds relationships that extend beyond the project timeline. Schools will serve as both consultation partners and audience development channels, integrating the murals into educational programming on history, visual arts, and community identity.')
doc.add_paragraph('The project\u2019s mentorship model creates collaborative relationships between established artists and emerging practitioners that are designed to outlast the project. These mentor-mentee connections become part of a growing network of creative professionals in the Nelson Mandela Bay region, fostering future collaborations on independent projects, joint commissions, and collective exhibitions.')

doc.add_heading('Strategy for Attracting New Partners', level=3)
doc.add_paragraph('APBC will leverage the Kariega Champs project as a demonstrable case study to attract new partnerships for future phases. The project\u2019s documented outcomes \u2014 employment data, community engagement records, media coverage, and visual documentation \u2014 will form a portfolio presented to corporate sponsors, tourism bodies, and cultural exchange programmes. The mural trail concept offers specific partnership opportunities for Nelson Mandela Bay tourism operators, heritage organisations, and local businesses seeking community investment vehicles.')

doc.add_heading('Long-term Vision', level=3)
doc.add_paragraph('APBC envisions the Kariega Champs project as the first phase of an ongoing public art programme that will expand to celebrate additional Kariega cultural figures and heritage themes. Each new phase creates partnership opportunities, allowing existing partners to deepen their involvement and new partners to join a proven initiative. This phased approach builds a growing ecosystem of shared resources, expertise, and funding that strengthens the sustainability of Kariega\u2019s public art infrastructure and APBC\u2019s role as a credible arts delivery organisation in the Eastern Cape.')

# SAVE
output = 'C:/Users/sinet/OneDrive/Documents/MyBliss Projects/2026 Projects/Working Files/Kariega_Champs_NAC_FINAL_v2.docx'
doc.save(output)
print(f'SUCCESS: Saved to {output}')
